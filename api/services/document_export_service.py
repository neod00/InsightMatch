"""
ISO Manual Document Export Service
Converts Markdown ISO manuals to PDF and Word (DOCX) formats.
"""

import io
import re
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, cm
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable, ListFlowable, ListItem
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


# ============================================================
# Korean Font Registration for PDF
# ============================================================
_FONT_REGISTERED = False
# 등록에 성공한 실제 한글 폰트 이름 (플랫폼에 따라 런타임에 결정됨).
# 나머지 코드는 이 두 전역을 참조하므로 'MalgunGothic' 하드코딩에 의존하지 않는다.
_KR_FONT = 'Helvetica'
_KR_FONT_BOLD = 'Helvetica-Bold'


def _register_korean_fonts():
    """한글 폰트를 등록한다. 플랫폼별로 아래 순서로 시도한다:
      1) 시스템/번들 TTF (Windows 맑은고딕, Linux 나눔고딕, 저장소 번들 폰트)
      2) ReportLab 내장 한국어 CID 폰트 (폰트 파일 불필요 — Vercel(Linux)에서 사용)
      3) Helvetica (한글 미표시, 단 최소한 크래시는 방지)
    성공 시 전역 _KR_FONT / _KR_FONT_BOLD 에 실제 등록 폰트명을 세팅한다.
    """
    global _FONT_REGISTERED, _KR_FONT, _KR_FONT_BOLD
    if _FONT_REGISTERED:
        return

    bundled_dir = os.path.join(os.path.dirname(__file__), 'assets', 'fonts')
    ttf_candidates = [
        # (normal_path, bold_path)
        (r'C:\Windows\Fonts\malgun.ttf', r'C:\Windows\Fonts\malgunbd.ttf'),
        ('/usr/share/fonts/truetype/nanum/NanumGothic.ttf',
         '/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf'),
        ('/usr/share/fonts/truetype/nanum/NanumGothic.ttf',
         '/usr/share/fonts/truetype/nanum/NanumGothic.ttf'),
        (os.path.join(bundled_dir, 'NanumGothic.ttf'),
         os.path.join(bundled_dir, 'NanumGothicBold.ttf')),
    ]

    for normal_path, bold_path in ttf_candidates:
        if normal_path and os.path.exists(normal_path):
            try:
                pdfmetrics.registerFont(TTFont('MalgunGothic', normal_path))
                bold_name = 'MalgunGothic'
                if bold_path and os.path.exists(bold_path):
                    pdfmetrics.registerFont(TTFont('MalgunGothicBold', bold_path))
                    bold_name = 'MalgunGothicBold'
                from reportlab.pdfbase.pdfmetrics import registerFontFamily
                registerFontFamily(
                    'MalgunGothic',
                    normal='MalgunGothic',
                    bold=bold_name,
                    italic='MalgunGothic',
                    boldItalic=bold_name,
                )
                _KR_FONT = 'MalgunGothic'
                _KR_FONT_BOLD = bold_name
                _FONT_REGISTERED = True
                print(f'[PDF] Korean TTF font registered: {normal_path}')
                return
            except Exception as e:
                print(f'[PDF] TTF registration failed ({normal_path}): {e}')

    # 2) ReportLab 내장 한국어 CID 폰트 — 폰트 파일이 없는 리눅스(Vercel) 환경 대응
    try:
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        pdfmetrics.registerFont(UnicodeCIDFont('HYSMyeongJo-Medium'))
        pdfmetrics.registerFont(UnicodeCIDFont('HYGothic-Medium'))
        _KR_FONT = 'HYSMyeongJo-Medium'
        _KR_FONT_BOLD = 'HYGothic-Medium'
        _FONT_REGISTERED = True
        print('[PDF] Korean CID font (HYGothic / HYSMyeongJo) registered')
        return
    except Exception as e:
        print(f'[PDF] CID font fallback failed: {e}')

    # 3) 최후의 폴백: Helvetica (한글은 표시되지 않지만 렌더링은 성공)
    _KR_FONT = 'Helvetica'
    _KR_FONT_BOLD = 'Helvetica-Bold'
    _FONT_REGISTERED = True
    print('[PDF] WARNING: No Korean font available — falling back to Helvetica (한글 미표시)')


# ============================================================
# Markdown Parser (shared between PDF and DOCX)
# ============================================================

def parse_markdown_lines(markdown_text):
    """
    Parse markdown into structured blocks.
    Returns list of dicts: {type, content, level, items, rows, ...}
    """
    blocks = []
    lines = markdown_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # --- Heading ---
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            blocks.append({
                'type': 'heading',
                'level': level,
                'content': heading_match.group(2).strip()
            })
            i += 1
            continue

        # --- Horizontal rule ---
        if stripped in ('---', '***', '___'):
            blocks.append({'type': 'hr'})
            i += 1
            continue

        # --- Table ---
        if '|' in stripped and stripped.startswith('|'):
            table_rows = []
            while i < len(lines) and '|' in lines[i].strip() and lines[i].strip().startswith('|'):
                row_text = lines[i].strip()
                # Skip separator rows (|---|---|)
                if re.match(r'^\|[\s\-:|]+\|$', row_text):
                    i += 1
                    continue
                cells = [c.strip() for c in row_text.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1
            if table_rows:
                blocks.append({'type': 'table', 'rows': table_rows})
            continue

        # --- Ordered list ---
        ol_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if ol_match:
            items = []
            while i < len(lines):
                ol_m = re.match(r'^\s*\d+\.\s+(.+)$', lines[i].strip())
                if ol_m:
                    items.append(ol_m.group(1))
                    i += 1
                else:
                    break
            blocks.append({'type': 'ordered_list', 'items': items})
            continue

        # --- Unordered list ---
        ul_match = re.match(r'^[-*•]\s+(.+)$', stripped)
        if ul_match:
            items = []
            while i < len(lines):
                ul_m = re.match(r'^\s*[-*•]\s+(.+)$', lines[i].strip())
                if ul_m:
                    items.append(ul_m.group(1))
                    i += 1
                else:
                    break
            blocks.append({'type': 'unordered_list', 'items': items})
            continue

        # --- Paragraph ---
        para_lines = []
        while i < len(lines):
            l = lines[i].strip()
            if not l:
                i += 1
                break
            if re.match(r'^#{1,4}\s+', l):
                break
            if l in ('---', '***', '___'):
                break
            if l.startswith('|') and '|' in l:
                break
            if re.match(r'^\d+\.\s+', l):
                break
            if re.match(r'^[-*•]\s+', l):
                break
            para_lines.append(l)
            i += 1

        if para_lines:
            text = ' '.join(para_lines)
            blocks.append({'type': 'paragraph', 'content': text})

    return blocks


def clean_inline_markdown(text):
    """Remove or convert inline markdown markers like **bold**, *italic*, etc."""
    # Bold: **text** → text
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Italic: *text* → text
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Code: `text` → text
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def rich_inline_markdown(text):
    """Convert inline markdown to ReportLab XML tags."""
    # Bold: **text** → <b>text</b>
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    # Italic: *text* → <i>text</i>
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    # Code: `text` → <font face="Courier">text</font>
    text = re.sub(r'`(.+?)`', r'<font face="Courier" size="8">\1</font>', text)
    # Escape special XML chars (but not our tags)
    text = text.replace('&', '&amp;')
    # Re-fix tags broken by &amp;
    # (no-op for now as our content shouldn't have & in tags)
    return text


# ============================================================
# PDF Export (ReportLab)
# ============================================================

def _get_pdf_styles():
    """Create custom styles for ISO manual PDF with Korean font support."""
    _register_korean_fonts()
    styles = getSampleStyleSheet()

    KR_FONT = _KR_FONT
    KR_FONT_BOLD = _KR_FONT_BOLD

    # Title style
    styles.add(ParagraphStyle(
        'ISOTitle',
        parent=styles['Title'],
        fontName=KR_FONT_BOLD,
        fontSize=20,
        spaceAfter=6 * mm,
        spaceBefore=0,
        textColor=colors.HexColor('#1A365D'),
        alignment=TA_CENTER,
    ))

    # H1
    styles.add(ParagraphStyle(
        'ISOH1',
        parent=styles['Heading1'],
        fontName=KR_FONT_BOLD,
        fontSize=16,
        spaceBefore=10 * mm,
        spaceAfter=4 * mm,
        textColor=colors.HexColor('#1A365D'),
        borderWidth=0,
        borderPadding=0,
        borderColor=colors.HexColor('#C5A880'),
    ))

    # H2
    styles.add(ParagraphStyle(
        'ISOH2',
        parent=styles['Heading2'],
        fontName=KR_FONT_BOLD,
        fontSize=13,
        spaceBefore=6 * mm,
        spaceAfter=3 * mm,
        textColor=colors.HexColor('#2D4A6F'),
    ))

    # H3
    styles.add(ParagraphStyle(
        'ISOH3',
        parent=styles['Heading3'],
        fontName=KR_FONT_BOLD,
        fontSize=11,
        spaceBefore=4 * mm,
        spaceAfter=2 * mm,
        textColor=colors.HexColor('#374151'),
    ))

    # H4
    styles.add(ParagraphStyle(
        'ISOH4',
        parent=styles['Heading4'],
        fontName=KR_FONT_BOLD,
        fontSize=10,
        spaceBefore=3 * mm,
        spaceAfter=2 * mm,
        textColor=colors.HexColor('#4B5563'),
    ))

    # Body text
    styles.add(ParagraphStyle(
        'ISOBody',
        parent=styles['Normal'],
        fontName=KR_FONT,
        fontSize=9.5,
        leading=15,
        spaceBefore=1 * mm,
        spaceAfter=2 * mm,
        alignment=TA_JUSTIFY,
        textColor=colors.HexColor('#1F2937'),
    ))

    # List item
    styles.add(ParagraphStyle(
        'ISOListItem',
        parent=styles['Normal'],
        fontName=KR_FONT,
        fontSize=9.5,
        leading=14,
        leftIndent=8 * mm,
        spaceBefore=0.5 * mm,
        spaceAfter=0.5 * mm,
        textColor=colors.HexColor('#374151'),
    ))

    # Table cell
    styles.add(ParagraphStyle(
        'ISOTableCell',
        parent=styles['Normal'],
        fontName=KR_FONT,
        fontSize=8.5,
        leading=12,
        textColor=colors.HexColor('#1F2937'),
    ))

    # Table header
    styles.add(ParagraphStyle(
        'ISOTableHeader',
        parent=styles['Normal'],
        fontName=KR_FONT_BOLD,
        fontSize=8.5,
        leading=12,
        textColor=colors.white,
    ))

    return styles


def _add_page_number(canvas, doc):
    """Add page number and footer to each page."""
    canvas.saveState()
    # Page number
    canvas.setFont(_KR_FONT, 8)
    canvas.setFillColor(colors.HexColor('#9CA3AF'))
    canvas.drawRightString(
        A4[0] - 20 * mm,
        15 * mm,
        f"Page {doc.page}"
    )
    # Footer line
    canvas.setStrokeColor(colors.HexColor('#E5E7EB'))
    canvas.setLineWidth(0.5)
    canvas.line(20 * mm, 18 * mm, A4[0] - 20 * mm, 18 * mm)
    # Footer text
    canvas.setFont(_KR_FONT, 7)
    canvas.drawString(20 * mm, 15 * mm, "InsightMatch AI ISO Manual Generator")
    canvas.restoreState()


def markdown_to_pdf(markdown_text, company_name="", target_iso=""):
    """
    Convert markdown text to a PDF byte stream.
    Returns: io.BytesIO containing the PDF
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=25 * mm,
        title=f"{company_name} {target_iso} 매뉴얼" if company_name else "ISO 매뉴얼",
        author="InsightMatch AI",
    )

    styles = _get_pdf_styles()
    story = []
    blocks = parse_markdown_lines(markdown_text)

    for block in blocks:
        btype = block['type']

        if btype == 'heading':
            level = block['level']
            text = rich_inline_markdown(block['content'])
            style_name = {1: 'ISOH1', 2: 'ISOH2', 3: 'ISOH3', 4: 'ISOH4'}.get(level, 'ISOH4')
            story.append(Paragraph(text, styles[style_name]))

            # Add line under H1
            if level == 1:
                story.append(HRFlowable(
                    width="100%",
                    thickness=1.5,
                    color=colors.HexColor('#C5A880'),
                    spaceAfter=3 * mm,
                ))

        elif btype == 'paragraph':
            text = rich_inline_markdown(block['content'])
            story.append(Paragraph(text, styles['ISOBody']))

        elif btype == 'hr':
            story.append(Spacer(1, 2 * mm))
            story.append(HRFlowable(
                width="100%",
                thickness=0.5,
                color=colors.HexColor('#D1D5DB'),
                spaceAfter=2 * mm,
            ))

        elif btype == 'unordered_list':
            for item in block['items']:
                text = rich_inline_markdown(item)
                story.append(Paragraph(f"• {text}", styles['ISOListItem']))

        elif btype == 'ordered_list':
            for idx, item in enumerate(block['items'], 1):
                text = rich_inline_markdown(item)
                story.append(Paragraph(f"{idx}. {text}", styles['ISOListItem']))

        elif btype == 'table':
            rows = block['rows']
            if not rows:
                continue

            # Build table data with Paragraphs
            table_data = []
            for row_idx, row in enumerate(rows):
                if row_idx == 0:
                    # Header row
                    table_data.append([
                        Paragraph(rich_inline_markdown(cell), styles['ISOTableHeader'])
                        for cell in row
                    ])
                else:
                    table_data.append([
                        Paragraph(rich_inline_markdown(cell), styles['ISOTableCell'])
                        for cell in row
                    ])

            # Calculate column widths
            num_cols = max(len(r) for r in table_data)
            available_width = A4[0] - 40 * mm
            col_width = available_width / num_cols

            # Normalize row lengths
            for row in table_data:
                while len(row) < num_cols:
                    row.append(Paragraph('', styles['ISOTableCell']))

            table = Table(table_data, colWidths=[col_width] * num_cols)
            table.setStyle(TableStyle([
                # Header styling
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1A365D')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), _KR_FONT_BOLD),
                ('FONTSIZE', (0, 0), (-1, 0), 8.5),
                # Body styling
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#1F2937')),
                # Alternating rows
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F9FAFB')]),
                # Grid
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D1D5DB')),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(Spacer(1, 2 * mm))
            story.append(table)
            story.append(Spacer(1, 2 * mm))

    # Build PDF
    doc.build(story, onFirstPage=_add_page_number, onLaterPages=_add_page_number)
    buffer.seek(0)
    return buffer


# ============================================================
# Word (DOCX) Export (python-docx)
# ============================================================

def _setup_docx_styles(document):
    """Configure custom styles for the Word document."""
    # Default font
    style = document.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.5

    # Heading 1
    h1 = document.styles['Heading 1']
    h1.font.size = Pt(18)
    h1.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)

    # Heading 2
    h2 = document.styles['Heading 2']
    h2.font.size = Pt(14)
    h2.font.color.rgb = RGBColor(0x2D, 0x4A, 0x6F)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = document.styles['Heading 3']
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    # Heading 4
    h4 = document.styles['Heading 4']
    h4.font.size = Pt(11)
    h4.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    h4.font.bold = True
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(4)


def _add_inline_formatting(paragraph, text):
    """Add runs with bold/italic formatting based on markdown markers."""
    # Split by **bold** and *italic* patterns
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x6B, 0x21, 0xA8)
        else:
            paragraph.add_run(part)


def markdown_to_docx(markdown_text, company_name="", target_iso=""):
    """
    Convert markdown text to a DOCX byte stream.
    Returns: io.BytesIO containing the DOCX
    """
    doc = Document()
    _setup_docx_styles(doc)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Add footer with page numbers
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    # Add page numbering field
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    run._element.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)
    footer_para.add_run(' / ')
    run2 = footer_para.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    run2._element.append(fldChar3)
    instrText2 = OxmlElement('w:instrText')
    instrText2.text = 'NUMPAGES'
    run2._element.append(instrText2)
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run2._element.append(fldChar4)

    # Parse and build
    blocks = parse_markdown_lines(markdown_text)

    for block in blocks:
        btype = block['type']

        if btype == 'heading':
            level = min(block['level'], 4)
            heading = doc.add_heading(level=level)
            _add_inline_formatting(heading, block['content'])

        elif btype == 'paragraph':
            para = doc.add_paragraph()
            _add_inline_formatting(para, block['content'])

        elif btype == 'hr':
            # Add a thin horizontal line
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            pPr = para._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'D1D5DB')
            pBdr.append(bottom)
            pPr.append(pBdr)

        elif btype == 'unordered_list':
            for item in block['items']:
                para = doc.add_paragraph(style='List Bullet')
                _add_inline_formatting(para, item)

        elif btype == 'ordered_list':
            for item in block['items']:
                para = doc.add_paragraph(style='List Number')
                _add_inline_formatting(para, item)

        elif btype == 'table':
            rows_data = block['rows']
            if not rows_data:
                continue

            num_cols = max(len(r) for r in rows_data)
            num_rows = len(rows_data)

            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for row_idx, row_data in enumerate(rows_data):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < num_cols:
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = ''
                        para = cell.paragraphs[0]
                        _add_inline_formatting(para, cell_text)
                        para.paragraph_format.space_after = Pt(2)

                        # Style header row
                        if row_idx == 0:
                            from docx.oxml.ns import nsdecls
                            from docx.oxml import parse_xml
                            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A365D"/>')
                            cell._element.get_or_add_tcPr().append(shading)
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.font.bold = True
                                run.font.size = Pt(9)

            # Add space after table
            doc.add_paragraph()

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
